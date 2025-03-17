import os
import json
import shutil
from pathlib import Path
from ollama import Client
from PyPDF2 import PdfReader
from docx import Document
from tqdm import tqdm

class DocumentConverter:
    def __init__(self):
        self.model = "llama3"  # Using llama2 13B as llama3 is not yet available in Ollama
        self.client = Client()
        
    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF file"""
        reader = PdfReader(file_path)
        text = ""
        for page in tqdm(reader.pages, desc="Processing PDF pages", unit="page"):
            text += page.extract_text() + "\n"
        return text

    def extract_table_data(self, table):
        """Extract data from a table"""
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Get text from the cell, removing extra whitespace
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            data.append(row_data)
        return data

    def extract_text_from_doc(self, file_path):
        """Extract text and tables from DOC/DOCX file"""
        doc = Document(file_path)
        content = {
            "paragraphs": [],
            "tables": []
        }

        # Extract paragraphs
        for paragraph in tqdm(doc.paragraphs, desc="Processing paragraphs", unit="para"):
            if paragraph.text.strip():  # Only add non-empty paragraphs
                content["paragraphs"].append(paragraph.text.strip())

        # Extract tables
        for table in tqdm(doc.tables, desc="Processing tables", unit="table"):
            table_data = self.extract_table_data(table)
            if table_data:  # Only add non-empty tables
                content["tables"].append(table_data)

        # Convert the content to a formatted string
        text = "Document Content:\n\n"
        
        # Add paragraphs
        if content["paragraphs"]:
            text += "Paragraphs:\n"
            for i, para in enumerate(content["paragraphs"], 1):
                text += f"Paragraph {i}: {para}\n"
            text += "\n"
        
        # Add tables
        if content["tables"]:
            text += "Tables:\n"
            for i, table in enumerate(content["tables"], 1):
                text += f"Table {i}:\n"
                for row in table:
                    text += f"| {' | '.join(row)} |\n"
                text += "\n"

        return text

    def process_text_with_ollama(self, text):
        """Process text using Ollama's Llama2 model"""
        try:
            # Create messages for the conversation
            messages = [
                {
                    "role": "system",
                    "content": "You are a helpful assistant that converts text into well-structured JSON. Always ensure the output is valid JSON format. When processing tables, maintain their structure in the JSON output."
                },
                {
                    "role": "user",
                    "content": f"Convert the following text into a structured JSON format, maintaining all information and context, especially preserving table structures. Text: {text}"
                }
            ]
            
            # Generate response using Ollama
            response = self.client.chat(model=self.model, messages=messages)
            
            if response and 'message' in response and 'content' in response['message']:
                return response['message']['content']
            return None
            
        except Exception as e:
            print(f"Error calling Ollama: {e}")
            return None

    def process_file(self, input_file):
        """Process a single file and convert it to JSON"""
        file_path = Path(input_file)
        
        if not file_path.exists():
            print(f"File not found: {input_file}")
            return
        
        # Extract text based on file type
        try:
            if file_path.suffix.lower() == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_path.suffix.lower() in ['.doc', '.docx']:
                text = self.extract_text_from_doc(file_path)
            else:
                print(f"Unsupported file type: {file_path.suffix}")
                return
            
            # Process text with Ollama
            json_content = self.process_text_with_ollama(text)
            if not json_content:
                print(f"Failed to process text with Ollama for file: {input_file}")
                return
            
            # Create output directory if it doesn't exist
            output_dir = Path('output')
            output_dir.mkdir(exist_ok=True)
            
            # Save JSON file
            json_file_path = output_dir / f"{file_path.stem}.json"
            with open(json_file_path, 'w', encoding='utf-8') as f:
                # Try to parse and pretty print the JSON if possible
                try:
                    json_obj = json.loads(json_content)
                    json.dump(json_obj, f, indent=2, ensure_ascii=False)
                except json.JSONDecodeError:
                    # If parsing fails, write the raw content
                    f.write(json_content)
            
            # Copy original file to output directory
            shutil.copy2(file_path, output_dir / file_path.name)
            
            # Remove the file from input directory after successful processing and copying
            file_path.unlink()
            
            #print(f"Successfully processed {input_file}")
            #print(f"JSON output saved to: {json_file_path}")
            #print(f"Original file moved to output directory and removed from input")
            
        except Exception as e:
            print(f"Error processing file {input_file}: {e}")

def main():
    converter = DocumentConverter()
    input_dir = Path('input')
    
    if not input_dir.exists():
        print("Input directory not found!")
        return
    
    # Get list of files to process
    files_to_process = [f for f in input_dir.glob('*.*') if f.suffix.lower() in ['.pdf', '.doc', '.docx']]
    
    # Process all PDF and DOC files in the input directory with progress bar
    for file_path in tqdm(files_to_process, desc="Processing files", unit="file"):
        converter.process_file(str(file_path))

if __name__ == "__main__":
    main() 